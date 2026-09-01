# ============================================================
# resume_parser.py
# PDF / DOCX resume text extraction
# ============================================================

import re

import PyPDF2
from docx import Document


# ============================================================
# CLEAN EXTRACTED TEXT
# ============================================================

def clean_extracted_text(text):
    """
    Clean text extracted from PDF/DOCX without destroying
    useful resume information.
    """

    if not text:
        return ""

    text = text.replace("\x00", " ")

    # Normalize line endings
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Remove excessive spaces
    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    # Remove excessive blank lines
    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# ============================================================
# PDF
# ============================================================

def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF resume.
    """

    text_parts = []

    with open(
        file_path,
        "rb"
    ) as file:

        reader = PyPDF2.PdfReader(
            file
        )

        for page in reader.pages:

            try:
                page_text = page.extract_text()
            except Exception:
                page_text = None

            if page_text:
                text_parts.append(
                    page_text
                )

    return clean_extracted_text(
        "\n".join(text_parts)
    )


# ============================================================
# DOCX
# ============================================================

def extract_text_from_docx(file_path):
    """
    Extract text from paragraphs and tables in a DOCX resume.
    """

    document = Document(
        file_path
    )

    text_parts = []

    # Paragraphs
    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text_parts.append(
                paragraph.text
            )

    # Tables
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(
                        cell_text
                    )

            if row_text:
                text_parts.append(
                    " | ".join(row_text)
                )

    return clean_extracted_text(
        "\n".join(text_parts)
    )


# ============================================================
# MAIN FUNCTION
# ============================================================

def extract_resume_text(file_path):
    """
    Detect file type and extract resume text.
    """

    lower_path = file_path.lower()

    if lower_path.endswith(".pdf"):

        return extract_text_from_pdf(
            file_path
        )

    if lower_path.endswith(".docx"):

        return extract_text_from_docx(
            file_path
        )

    raise ValueError(
        "Only PDF and DOCX files are supported."
    )